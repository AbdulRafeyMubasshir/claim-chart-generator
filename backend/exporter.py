"""
exporter.py

Generates a properly formatted claim chart as a .docx file.
Patent attorneys work in Word — this is a practical requirement,
not just a nice-to-have.

Claim chart format:
| Claim Element | Evidence from Product Document | Notes |
"""

import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from patent_parser import Claim
from mapper import MappingResult


def generate_claim_chart(
    claim: Claim,
    mappings: list[MappingResult],
    patent_title: str,
    product_name: str
) -> bytes:
    """
    Generate a claim chart Word document.

    Returns the document as bytes (for HTTP response).
    """
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    title = doc.add_heading("Claim Chart", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f"Patent: {patent_title}")
    doc.add_paragraph(f"Product / Document: {product_name}")
    doc.add_paragraph(f"Claim {claim.number} ({'Independent' if claim.is_independent else f'Dependent on Claim {claim.parent_claim}'})")
    doc.add_paragraph("")

    # Claim text box
    claim_para = doc.add_paragraph()
    claim_para.add_run("Full Claim Text:").bold = True
    doc.add_paragraph(claim.text)
    doc.add_paragraph("")

    # Build the table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Set column widths
    widths = [Inches(2.8), Inches(4.5), Inches(1.7)]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = width

    # Header row
    header_cells = table.rows[0].cells
    headers = ["Claim Element / Limitation", "Evidence from Product Document", "Notes"]
    for cell, header_text in zip(header_cells, headers):
        cell.text = header_text
        # Bold headers
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
        _set_cell_background(cell, "1F3864")  # Dark navy
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for i, mapping in enumerate(mappings):
        row = table.add_row()
        cells = row.cells

        # Column 1: Claim element
        cells[0].text = f"Element {mapping.claim_element.index + 1}:\n{mapping.claim_element.text}"

        # Column 2: Evidence
        if mapping.found_evidence:
            evidence_para = cells[1].paragraphs[0]
            evidence_para.add_run(mapping.evidence_passage)
            # Add explanation below
            cells[1].add_paragraph("")
            exp_para = cells[1].add_paragraph()
            exp_run = exp_para.add_run(f"Mapping: {mapping.explanation}")
            exp_run.font.size = Pt(9)
            exp_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        else:
            no_ev = cells[1].paragraphs[0]
            no_ev_run = no_ev.add_run("No clear evidence found in provided document.")
            no_ev_run.font.color.rgb = RGBColor(0xAA, 0x44, 0x44)
            if mapping.explanation:
                cells[1].add_paragraph(mapping.explanation)

        # Column 3: Notes / Confidence
        confidence_text = mapping.confidence.upper() if mapping.found_evidence else "N/A"
        cells[2].text = f"Confidence: {confidence_text}"
        if not mapping.found_evidence:
            cells[2].add_paragraph("Manual review needed")

        # Alternating row shading
        if i % 2 == 0:
            for cell in cells:
                _set_cell_background(cell, "EBF0FA")

    # Footer note
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Note: This claim chart was generated with AI assistance. "
        "All mappings should be reviewed by a qualified patent attorney "
        "before use in any legal proceeding."
    )
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True

    # Serialize to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _set_cell_background(cell, hex_color: str):
    """Set the background colour of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)
